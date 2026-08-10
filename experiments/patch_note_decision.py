"""
patch_note_decision.py -- Point 0 du BRIEF "NsDiff : rouvrir la question
economique par le rapport edge/frais" : hygiene documentaire de la
`Note_decision_DeepEdgeBenchmark.docx`.

TROIS CORRECTIONS, et rien d'autre :

  1. ETAT DU DASHBOARD -- la note affirme encore que « le dashboard D7/W1 tourne
     en single-seed/50, scripts prets et smoke-testes, non executes ». C'est FAUX
     depuis le 6/8 : `repoint_oos_to_m200.json` porte `applied=true`,
     run `20260806-oos-repoint-m200`, sauvegarde
     `tracking.db.bak_repoint_m200_2026-08-06T064808`, couverture 0,909 -> 0,9315
     sur 2 700 lignes. Le dashboard est en graine 42 x 200 tirages. Ce qui reste
     ouvert, et seulement cela : l'integration de l'ENSEMBLE 5x200 au dashboard
     (la piste `oos` porte une graine, pas un melange de cinq).

  2. `effective_n` MENSUEL, 12 vs 13 -- ce n'est pas une contradiction mais deux
     protocoles differents, que ni l'un ni l'autre document ne rendait explicite :
        * 40 origines de test, blocs de 3 -> effective_n = 13 (piste
          daily-pousse-vs-natif, `NOTE_compare_daily_vs_monthly_nsdiff`) ;
        * 36 origines de test, blocs de 3 -> effective_n = 12 (etude de
          faisabilite du chantier C, `monthly_feasibility.json`).
     La correction rend le compte d'origines explicite des deux cotes, pour que
     les deux chiffres cessent de se lire comme un desaccord.

  3. RESULTATS MANQUANTS -- la note ne mentionne ni le chantier B3 (pricing
     d'option et fonctionnelles de trajectoire) ni A3-ii (cadence de refit).
     Tous deux renforcent la recommandation « NO-GO trading » en fermant deux
     explications alternatives ; les omettre affaiblit la note en soutenance.

DISCIPLINE, identique a `repoint_oos_to_m200.py` : dry-run par defaut,
sauvegarde horodatee avant toute ecriture, `--apply` explicite. Le patch est
IDEMPOTENT : il verifie le texte attendu avant de le remplacer et refuse de
tourner deux fois (il signale que la correction est deja en place).

La mise en forme est preservee : chaque paragraphe est reecrit run par run a
partir du formatage des runs existants (gras / italique), jamais aplati.

Sortie : experiments/patch_note_decision.json (journal des remplacements)
Usage :
    python patch_note_decision.py                 # dry-run : montre les diffs
    python patch_note_decision.py --apply         # sauvegarde + patch
"""

import argparse
import json
import shutil
import sys
import time
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_DOC = Path.home() / "Downloads" / "Note_decision_DeepEdgeBenchmark.docx"
OUT_PATH = Path(__file__).resolve().parent / "patch_note_decision.json"


# ── reecriture d'un paragraphe en preservant le formatage ───────────────────

def set_runs(paragraph, segments) -> None:
    """Reecrit le paragraphe avec `segments` = [(texte, gras, italique), ...].

    Les runs existants sont reutilises tant qu'il y en a (leur formatage de
    police -- taille, couleur, style de caractere -- est ainsi conserve) ; le
    gras et l'italique sont poses explicitement. Les runs en trop sont vides,
    les manquants ajoutes en copiant le dernier."""
    runs = paragraph.runs
    while len(runs) < len(segments):
        paragraph.add_run("")
        runs = paragraph.runs
    for i, (text, bold, italic) in enumerate(segments):
        runs[i].text = text
        runs[i].bold = bold
        runs[i].italic = italic
    for r in runs[len(segments):]:
        r.text = ""
        r.bold = None
        r.italic = None


def insert_paragraph_after(paragraph, segments):
    """Nouveau paragraphe juste apres `paragraph`, meme style. python-docx
    n'expose pas l'insertion : on passe par l'arbre XML, ce qui est le chemin
    officiel documente pour cette operation."""
    import copy
    new_p = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, paragraph._parent)
    set_runs(para, segments)
    return para


# ── les corrections, declarees ──────────────────────────────────────────────

NEW_P54 = [
    ("Réserves et suite. ", True, None),
    ("État réel au 7/8 : la bascule 50→200 ", None, None),
    ("a été exécutée", True, None),
    (" le 6/8 (run 20260806-oos-repoint-m200, 2 700 lignes NsDiff, sauvegarde "
     "tracking.db.bak_repoint_m200_2026-08-06T064808, couverture 0,909 → 0,9315). "
     "Le dashboard D7/W1 tourne donc en ", None, None),
    ("graine 42 × 200 tirages", True, None),
    (", et non plus en single-seed/50. Une seule chose reste ouverte : "
     "l'intégration au dashboard de l'", None, None),
    ("ensemble 5×200", True, None),
    (" — la piste oos porte une graine, pas un mélange de cinq, si bien que la "
     "configuration de production n'est pas encore celle qu'affiche la page. "
     "Si l'on veut vraiment étayer « optimal », un balayage du budget de tirages "
     "est le seul test qui le prouverait — lancé depuis (chantier 3 du brief "
     "edge/frais).", None, None),
]

NEW_T_CADRAGE = ("« Suffisant et équitable », pas « optimal » prouvé ; bascule 200 exécutée le 6/8, "
                 "reste l'intégration de l'ensemble 5×200 au dashboard")

NEW_P62_TAIL = (" par actif sur cette piste, et comme M+1/M+2/M+3 se chevauchent, ")
NEW_P62_EFFN = "effective_n = 40 // 3 = 13"
NEW_P62_AFTER = (" par cellule (contre ~30 au weekly). L'étude de faisabilité du chantier C, "
                 "elle, teste 36 origines et affiche donc effective_n = 36 // 3 = 12 : ce n'est pas "
                 "un désaccord entre les deux documents mais deux grilles de test différentes, "
                 "même formule et mêmes blocs de 3. À cette échelle, « indistinguable » ne veut "
                 "pas dire « équivalent » : c'est l'absence de preuve du contraire, faute de "
                 "puissance.")

NEW_T_MENSUEL = ("~130 mois → effective_n = 13 (40 origines de test ; 12 sur la grille à 36 "
                 "origines du chantier C) ; mensuel-natif sur-couvre (PI 2–4× trop larges) ; "
                 "le critère d'arrêt du brief est atteint")

NEW_P_B3_A3 = [
    ("Deux explications alternatives ont été fermées. ", True, None),
    ("D'abord les ", None, None),
    ("scénarios de trajectoire", True, None),
    (" (chantier B3) : sur trois fonctionnelles qu'un couple de quantiles marginaux ne "
     "donne pas — minimum de parcours sur 3 semaines, put ATM, digital à barrière —, à "
     "budget strictement égal de 1 000 trajectoires de chaque côté, les deux modèles sont "
     "indistinguables sur les deux fonctionnelles franchement dépendantes du chemin, et sur "
     "le pricing du put ", None, None),
    ("GARCH gagne 7 cellules sur 10", True, None),
    (". La structure jointe de la diffusion n'apporte donc rien de mesurable, et là où une "
     "différence par cellule est nette, elle joue contre NsDiff. Ensuite la ", None, None),
    ("fraîcheur du fit", True, None),
    (" (chantier A3-ii) : refit trimestriel (×7,4 en coût de calcul) ou mensuel (", None, None),
    ("×24,6", True, None),
    ("), le verdict poolé contre GARCH est identique dans les trois bras — y compris sur la "
     "seule cellule que GARCH gagne (W+1 daily). L'asymétrie de protocole « GARCH est refit "
     "à chaque origine, NsDiff une seule fois » ne peut donc plus être invoquée comme "
     "explication de l'écart.", None, None),
]


def paragraph_text(segments) -> str:
    return "".join(s[0] for s in segments)


def find_paragraph(doc, *needles):
    """Repere un paragraphe par son CONTENU, jamais par son index. Necessaire :
    la correction 3 insere un paragraphe, ce qui decale tous les index suivants
    -- un reperage positionnel rendrait le script non rejouable (constate au
    premier run : le patch se croyait devant un paragraphe inattendu et
    refusait de tourner)."""
    for p in doc.paragraphs:
        if all(n in p.text for n in needles):
            return p
    return None


def main() -> None:
    p = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    p.add_argument("--doc", default=str(DEFAULT_DOC))
    p.add_argument("--apply", action="store_true", help="écrit réellement (défaut : dry-run)")
    p.add_argument("--out", default=str(OUT_PATH))
    args = p.parse_args()

    doc_path = Path(args.doc)
    if not doc_path.exists():
        raise SystemExit(f"document introuvable : {doc_path}")
    d = Document(str(doc_path))

    changes, already = [], []

    # ── 1. etat du dashboard ────────────────────────────────────────────────
    # sentinelle : la phrase PERIMEE, pas le simple mot "single-seed/50" -- le
    # texte corrige le cite lui aussi ("et non plus en single-seed/50"), et une
    # sentinelle trop large rendait le patch non idempotent.
    stale = find_paragraph(d, "Réserves et suite", "tourne encore en single-seed/50")
    if stale is not None:
        changes.append({"cible": "§3 « Réserves et suite »",
                        "motif": "affirme à tort que la bascule 200 n'est pas exécutée",
                        "avant": stale.text, "apres": paragraph_text(NEW_P54)})
        if args.apply:
            set_runs(stale, NEW_P54)
    elif find_paragraph(d, "Réserves et suite", "la bascule 50→200 a été exécutée") is not None:
        already.append("§3 « Réserves et suite » déjà corrigé")
    else:
        raise SystemExit("§3 « Réserves et suite » introuvable dans un état connu -- patch refusé")

    cell = d.tables[0].rows[3].cells[3]
    if "encore à exécuter" in cell.text:
        changes.append({"cible": "tableau de synthèse, ligne « Cadrage 200 tirages », réserve",
                        "motif": "même affirmation, dans le tableau",
                        "avant": cell.text, "apres": NEW_T_CADRAGE})
        if args.apply:
            set_runs(cell.paragraphs[0], [(NEW_T_CADRAGE, False, None)])
    else:
        already.append("cellule « Cadrage 200 tirages » déjà corrigée")

    # ── 2. effective_n mensuel ──────────────────────────────────────────────
    p62 = find_paragraph(d, "contrainte de puissance qui domine tout")
    if p62 is None:
        raise SystemExit("§4 « contrainte de puissance » introuvable -- patch refusé")
    if "effective_n ≈ 13" in p62.text and "36 origines" not in p62.text:
        segs = [(r.text, r.bold, r.italic) for r in p62.runs]
        segs[5] = (NEW_P62_TAIL, None, None)
        segs[6] = (NEW_P62_EFFN, True, None)
        segs[7] = (NEW_P62_AFTER, None, None)
        changes.append({"cible": "§4 « une contrainte de puissance qui domine tout »",
                        "motif": "12 vs 13 sans explication -- expliciter les deux grilles",
                        "avant": p62.text, "apres": paragraph_text(segs)})
        if args.apply:
            set_runs(p62, segs)
    else:
        already.append("§4 « contrainte de puissance » déjà harmonisé")

    cell = d.tables[0].rows[4].cells[2]
    if "effective_n≈13" in cell.text:
        changes.append({"cible": "tableau de synthèse, ligne « Cadrage mensuel »",
                        "motif": "même harmonisation, dans le tableau",
                        "avant": cell.text, "apres": NEW_T_MENSUEL})
        if args.apply:
            set_runs(cell.paragraphs[0], [(NEW_T_MENSUEL, False, None)])
    else:
        already.append("cellule « Cadrage mensuel » déjà harmonisée")

    # ── 3. resultats B3 et A3-ii ────────────────────────────────────────────
    anchor = find_paragraph(d, "le face-à-face direct", "24 tests appariés")
    if anchor is None:
        raise SystemExit("§2 « face-à-face direct » introuvable -- patch refusé")
    if find_paragraph(d, "chantier B3") is None:
        changes.append({"cible": "nouveau paragraphe après « le face-à-face direct » (§2)",
                        "motif": "B3 (pricing d'option, GARCH 7/10) et A3-ii (refit ×24,6) absents",
                        "avant": "(aucun)", "apres": paragraph_text(NEW_P_B3_A3)})
        if args.apply:
            insert_paragraph_after(anchor, NEW_P_B3_A3)
    else:
        already.append("paragraphe B3/A3-ii déjà présent")

    report = {"document": str(doc_path), "applied": bool(args.apply),
              "n_changes": len(changes), "changes": changes, "already_correct": already}

    print(f"Document : {doc_path}")
    for c in changes:
        print(f"\n--- {c['cible']}\n    motif : {c['motif']}")
        print(f"    AVANT : {c['avant'][:260]}")
        print(f"    APRÈS : {c['apres'][:260]}")
    for a in already:
        print(f"  (déjà en place) {a}")

    if not args.apply:
        print(f"\n--dry-run : rien écrit. {len(changes)} correction(s) prête(s). Relancer avec --apply.")
    elif changes:
        backup = doc_path.with_suffix(f".bak_{time.strftime('%Y-%m-%dT%H%M%S')}.docx")
        shutil.copy2(doc_path, backup)
        report["backup"] = str(backup)
        d.save(str(doc_path))
        print(f"\nSauvegarde -> {backup.name}\nDocument patché : {len(changes)} correction(s).")
    else:
        print("\nRien à corriger : le document est déjà à jour.")

    Path(args.out).write_text(json.dumps(report, indent=2, ensure_ascii=False))
    print(f"Journal -> {args.out}")


if __name__ == "__main__":
    main()
